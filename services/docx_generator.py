from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import io
import re

def create_docx_from_text(text: str):
    """Create a DOCX attempting to preserve simple resume template structure.

    Heuristics applied:
      - Blank lines preserved.
      - Lines starting with bullet markers (-, *, •) converted to bullet list items.
      - Uppercase short lines (likely section headings) bolded.
      - Normal lines become plain paragraphs.
    This does not replicate original fonts/margins from an uploaded DOCX/PDF; for
    pixel-perfect preservation we would need to diff and modify the original file.
    """
    document = Document()

    if not text:
        text = "(empty)"

    lines = text.splitlines()

    heading_pattern = re.compile(r"^[A-Z0-9 &/()'-]{3,}$")  # crude detection
    bullet_pattern = re.compile(r"^\s*([-*•])\s+(.*)$")

    for raw in lines:
        line = raw.rstrip('\r')
        if line.strip() == "":
            document.add_paragraph("")
            continue
        m_bullet = bullet_pattern.match(line)
        if m_bullet:
            content = m_bullet.group(2).strip()
            p = document.add_paragraph(style='List Bullet')
            run = p.add_run(content)
            run.font.size = Pt(10.5)
            continue
        # Heading heuristic
        if heading_pattern.match(line.strip()) and len(line.strip().split()) <= 5:
            p = document.add_paragraph()
            run = p.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(11)
            continue
        # Regular paragraph
        p = document.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10.5)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer